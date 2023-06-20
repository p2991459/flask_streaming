from docx import Document
from tabulate import tabulate

def read_table(table):
    table_content = []
    for row in table.rows:
        row_content = []
        for cell in row.cells:
            cell_text = cell.text.strip()  # Retrieve the text from the cell and remove leading/trailing spaces
            row_content.append(cell_text)
        table_content.append(row_content)
    return table_content
def tabulate_table(table_content):
    table_text = tabulate(table_content, headers="firstrow", tablefmt="pipe")
    return table_text

# Usage example
if __name__ == '__main__':
    input_file = 'SRS.docx'
    output_file = 'output.docx'
    doc = Document(input_file)
    # table_text = ''
    for table in doc.tables:
        table_data = read_table(table)
        table_text = tabulate_table(table_data)
        prompt = f"table_text: {table_text}"
        response = ask_a_question(prompt)
        collected_chunks = []
        collected_messages = []
        # iterate through the stream of events
        for chunk in response:
            chunk_time = time.time() - start_time  # calculate the time delay of the chunk
            collected_chunks.append(chunk)  # save the event response
            chunk_message = chunk['choices'][0]['delta']  # extract the message
            if "content" in chunk_message.keys():
                chunk_message = chunk['choices'][0]['delta']["content"]
            collected_messages.append(chunk_message)  # save the message
            print(f"{chunk_message}", end=" ")
# open("table_text.txt","w").write(table_text)


# for paragraph in doc.paragraphs:
#     print(paragraph.text)

# print(doc.element.body.xml) //To get the doc body as xml